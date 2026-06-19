"""Testes do export .docx (Fatia F2.1). Puro.

Aceite: o .docx gerado REABRE sem corrupcao (Document consegue le-lo) e tem a
estrutura esperada (capa, titulos com estilo Heading, secao de Notas, marcador
sobrescrito). E o template de estilo aceita overrides preservando o default.
"""

import io

from docx import Document

from open_webui.editorial.export.docx_export import build_docx
from open_webui.editorial.export.style import resolve_style
from open_webui.editorial.extractors.docx import extract_docx
from open_webui.test.editorial.test_docx_extractor import build_docx_with_footnote


def test_build_docx_reopens_without_corruption_and_has_structure():
    tree = extract_docx(build_docx_with_footnote())
    data = build_docx(tree)

    d = Document(io.BytesIO(data))  # reabre => arquivo valido (sem corrupcao)
    styles = [p.style.name for p in d.paragraphs]
    texts = [p.text for p in d.paragraphs]

    assert any(s == "Heading 1" and t == "Capitulo 1" for s, t in zip(styles, texts))
    assert "Notas" in texts
    assert any(t.startswith("1. Ver Documento Fundador") for t in texts)

    supers = [
        r.text for p in d.paragraphs for r in p.runs if r.font.superscript
    ]
    assert "1" in supers, "marcador de nota sobrescrito deveria estar no corpo"


def test_style_template_merge_keeps_defaults():
    s = resolve_style({"page": {"margin_mm": 40}, "fonts": {"body": "Georgia"}})
    assert s["page"]["margin_mm"] == 40
    assert s["fonts"]["body"] == "Georgia"
    assert s["colors"]["heading"] == "2E3A29"  # default Nidum preservado
    assert s["toc"]["show"] is True
