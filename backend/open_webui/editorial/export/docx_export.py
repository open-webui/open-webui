"""Export da Arvore de Blocos -> .docx (Fatia F2.1).

Gera capa, sumario AUTOMATICO (campo TOC do Word, que popula ao abrir/F9),
quebra de pagina por capitulo (heading nivel 1), listas, tabelas e secao de
notas no fim (com marcadores sobrescritos no corpo). Aplica o manual de estilo
Nidum (fontes, tamanhos, cor de titulo, margens). python-docx importado de forma
preguicosa. Sem imports de open_webui (alem do style, que e puro).
"""

import io

from open_webui.editorial.export.style import resolve_style


def _footnote_map(blocks):
    return {b["id"]: b.get("text", "") for b in blocks if b.get("type") == "footnote"}


def _page_break(doc):
    from docx.enum.text import WD_BREAK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_toc_field(doc, label):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = doc.add_paragraph().add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = label
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, txt, end):
        run._r.append(el)


def _render_inlines(para, block, fn_text, counter, notes_out):
    inlines = block.get("inlines") or [
        {"t": "text", "s": block.get("text", ""), "marks": []}
    ]
    for inl in inlines:
        if inl.get("t") == "text":
            run = para.add_run(inl.get("s", ""))
            marks = inl.get("marks") or []
            if "bold" in marks:
                run.bold = True
            if "italic" in marks:
                run.italic = True
            if "underline" in marks:
                run.underline = True
        elif inl.get("t") == "footnote_ref":
            counter += 1
            run = para.add_run(str(counter))
            run.font.superscript = True
            notes_out.append((counter, fn_text.get(inl.get("id"), "")))
    return counter


def _render_table(doc, block):
    rows = (block.get("table") or {}).get("rows") or []
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r in rows:
        cells = table.add_row().cells
        for i, c in enumerate(r):
            cells[i].text = c.get("text", "")


def build_docx(tree: dict, overrides: dict = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, RGBColor

    style = resolve_style(overrides)
    blocks = tree.get("blocks", [])
    fn_text = _footnote_map(blocks)
    title = (tree.get("metadata") or {}).get("title") or "Documento"

    doc = Document()

    sec = doc.sections[0]
    margin = Mm(style["page"]["margin_mm"])
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = margin
    if str(style["page"]["size"]).upper() == "A4":
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)

    normal = doc.styles["Normal"]
    normal.font.name = style["fonts"]["body"]
    normal.font.size = Pt(style["sizes"]["body_pt"])

    hcolor = RGBColor.from_string(style["colors"]["heading"])
    for lvl, key in [(1, "h1_pt"), (2, "h2_pt"), (3, "h3_pt"), (4, "h4_pt")]:
        try:
            st = doc.styles["Heading %d" % lvl]
            st.font.name = style["fonts"]["heading"]
            st.font.size = Pt(style["sizes"][key])
            st.font.color.rgb = hcolor
        except KeyError:
            pass

    # Capa
    if style["cover"]["show"]:
        for _ in range(8):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(32)
        r.font.color.rgb = hcolor
        if style["cover"].get("subtitle"):
            ps = doc.add_paragraph()
            ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ps.add_run(style["cover"]["subtitle"]).font.size = Pt(14)
        _page_break(doc)

    # Sumario automatico
    if style["toc"]["show"]:
        tp = doc.add_paragraph()
        tr = tp.add_run(style["toc"]["title"])
        tr.bold = True
        tr.font.size = Pt(style["sizes"]["h1_pt"])
        tr.font.color.rgb = hcolor
        _add_toc_field(doc, "Sumario - atualize com F9 no Word")
        _page_break(doc)

    # Corpo
    note_counter = 0
    notes_out = []
    started = False
    for b in blocks:
        t = b.get("type")
        if t == "footnote":
            continue
        if t == "heading":
            lvl = min(int(b.get("level", 1)), 4)
            if lvl == 1 and started:
                _page_break(doc)  # cada capitulo em pagina nova
            doc.add_heading(b.get("text", ""), level=lvl)
            started = True
        elif t == "list_item":
            para = doc.add_paragraph(style="List Bullet")
            note_counter = _render_inlines(para, b, fn_text, note_counter, notes_out)
            started = True
        elif t == "table":
            _render_table(doc, b)
            started = True
        else:
            para = doc.add_paragraph()
            note_counter = _render_inlines(para, b, fn_text, note_counter, notes_out)
            started = True

    # Notas no fim
    if notes_out:
        _page_break(doc)
        doc.add_heading(style["notes"]["title"], level=1)
        for num, text in notes_out:
            doc.add_paragraph("%d. %s" % (num, text))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
