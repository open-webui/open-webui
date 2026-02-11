from fastapi import APIRouter, HTTPException, Depends, Request
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Dict, List, Optional
import logging
import os
import json
from urllib.parse import urlparse
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


from open_webui.models.knowledge import Knowledges
from open_webui.retrieval.vector.connector import VECTOR_DB_CLIENT
from open_webui.utils.auth import get_verified_user
from open_webui.retrieval.web.tavily import search_tavily
from open_webui.retrieval.web.main import SearchResult

# facilities prompt template
FACILITIES_PROMPT = """You are an expert grant writer specializing in the 'Facilities, Equipment, and Other Resources' section of academic proposals for {sponsor} grants.

Your job is to expand and professionally refine the user's section draft using:
1.  User Input (this forms the base and must be preserved).
2.  PDF Research Chunks from relevant proposals or facility descriptions.
3.  Trusted Web Snippets


---

### FEW-SHOT GUIDANCE

**Example 1**  
> *User:* "Our Robotics Lab has 10 industrial arms and a motion capture system."  
> *Context:* `<source><source_id>Robotics_Lab_Equipment.pdf</source_id><source_context>Lab has 10 KUKA robotic arms and OptiTrack system</source_context></source>`
> *Draft:* "The Robotics Lab features 10 industrial-grade robotic manipulators and an advanced OptiTrack motion capture system for multi-agent interaction studies [Robotics_Lab_Equipment.pdf]."

**Example 2**  
> *User:* "We use our university's HPC system with 1000 GPUs."  
> *Context:* `<source><source_id>HPC_Cluster_Overview.pdf</source_id><source_context>University HPC cluster with 1,024 NVIDIA A100 GPUs</source_context></source>`
> *Draft:* "High-performance computing is supported by the university's HPC cluster with 1,024 NVIDIA A100 GPUs [HPC_Cluster_Overview.pdf]."

**Example 3 - Multiple Sources**  
> *User:* "Our wireless lab has software-defined radios and spectrum analyzers."  
> *Context:* `<source><source_id>Lab_Inventory_2023.pdf</source_id><source_context>12 USRP B210 software-defined radios</source_context></source>` and `<source><source_id>Equipment_Manual_2024.pdf</source_id><source_context>Keysight N9020A spectrum analyzers</source_context></source>`
> *Draft:* "The wireless communications laboratory is equipped with 12 USRP B210 software-defined radios [Lab_Inventory_2023.pdf] and Keysight N9020A spectrum analyzers [Equipment_Manual_2024.pdf] for advanced signal processing and spectrum analysis."
> *NOTE:* Use separate brackets [source1.pdf] [source2.pdf] NOT combined [source1.pdf; source2.pdf]

---

### SECTION: {section}
{section_specific_instructions}

**User Input:**
\"\"\"
{user_input}
\"\"\"

**PDF Research Chunks:**
\"\"\"
{retrieved_chunks}
\"\"\"

**Web Snippets**
\"\"\"
{web_snippets}
\"\"\"

---

### INSTRUCTIONS
- **IMPORTANT: You are GENERATING and REWRITING content, NOT copying text verbatim. Synthesize information from sources to create professional, grant-worthy prose.**
- Start with the User Input, retain all core ideas.
- ONLY use information from the User Input, PDF Chunks, and Web Snippets provided above. DO NOT add any information not present in these sources.
- Do not invent, assume, or add any equipment names, specifications, capabilities, or details not explicitly mentioned in the provided sources.
- Expand extensively with factual, cited data from PDF Chunks or Web Snippets, but only use information that is actually present in those sources.
- Use specific technical details, numbers, equipment names, capabilities, and concrete information ONLY if they are explicitly mentioned in the provided sources.
- Write comprehensive, detailed content that demonstrates expertise and thoroughness - aim for substantial, grant-worthy content using only the provided sources.
- Incorporate information from multiple PDF sources and web snippets to create rich, well-cited content, but only cite information that is actually present in those sources.
- Include specific equipment models, technical specifications, capabilities, and research applications ONLY if they are explicitly mentioned in the provided PDF Chunks or Web Snippets.
- Do not use any external knowledge, assumptions, or information not provided in the sources above.
- **When different sentences or parts of a paragraph come from different sources, cite each source appropriately. Example: "The lab has 10 robotic arms [source1.pdf] and advanced motion capture systems [source2.pdf] for multi-agent research."**
{integration_instructions}
**CRITICAL CITATION RULES - FOLLOW EXACTLY:**

**RULE 1: SEPARATE BRACKETS FOR EACH SOURCE (NEVER COMBINE!)**
- Each source MUST have its own separate bracket
- CORRECT: "...research facilities [source1.pdf] [https://nyu.edu/research] are available."
- WRONG: "...research facilities [source1.pdf; https://nyu.edu/research] are available."
- WRONG: "...research facilities [source1.pdf, https://nyu.edu/research] are available."
- If citing multiple sources for the same claim, use separate brackets with a space between them
- NEVER use semicolons (;) or commas (,) inside brackets to combine sources
- **CRITICAL: If the same information appears in multiple files (same content in different <source_id> tags), you MUST cite ALL of them. Example: "The HPC cluster has 1000 GPUs [panwar.pdf] [rangan.pdf]."**
- **CRITICAL: When different sentences or parts of a paragraph come from different sources, cite each source for its respective information. Example: "The lab includes 10 robotic arms [source1.pdf] and advanced motion capture systems [source2.pdf] for multi-agent research."**

**RULE 2: ONLY CITE SOURCES FROM <source_id> TAGS**
- **IMPORTANT: When web search is disabled, ONLY cite PDF files and document names - NEVER cite web URLs even if they appear in sources.**
- **ONLY use source names/links that appear in `<source_id>` tags in the context above.**
- **Look at the context above and find the exact `<source_id>` values.**
- **If you see `<source><source_id>HRamani_NSFCAREER_SUBMITTED_Proposal_20240723_NSF.pdf</source_id>`, you MUST use [HRamani_NSFCAREER_SUBMITTED_Proposal_20240723_NSF.pdf].**
- **NEVER use numbered citations like [1], [2], [3] - ALWAYS use the actual source name/link.**
- **Do NOT use source names/links that are NOT in the `<source_id>` tags.**
- **If no `<source_id>` tags are provided, do not include any citations.**
- **If web search is disabled, do NOT include any web URLs as citations, even if they appear in knowledge base sources.**
- Never invent or assume data not present in input or sources.
- If no relevant info is found, return only the user's input — improved stylistically.
- Focus on creating professional, grant-worthy content that showcases facilities and capabilities.
- **IMPORTANT: Do NOT add "Facilities, Equipment, and Other Resources" as a subtitle - just write the content for the section directly.**

Write a polished, comprehensive section suitable for direct inclusion in a {sponsor} grant.
"""

router = APIRouter()

class FacilitiesRequest(BaseModel):
    sponsor: str  
    form_data: Dict[str, str]  
    model: str  
    web_search_enabled: bool = False
    files: Optional[List[Dict]] = None

class FacilitiesResponse(BaseModel):
    success: bool
    message: str
    content: str 
    sources: List[Dict]  
    sections: Dict[str, str]  
    error: Optional[str] = None

class ExtractFormDataRequest(BaseModel):
    sponsor: str
    model: str
    files: List[Dict]

class ExtractFormDataResponse(BaseModel):
    success: bool
    form_data: Dict[str, str]
    error: Optional[str] = None

class SingleSectionRequest(BaseModel):
    sponsor: str
    section_key: str
    section_text: str
    model: str
    web_search_enabled: bool = False
    files: Optional[List[Dict]] = None

class SingleSectionResponse(BaseModel):
    success: bool
    section_key: str
    section_label: str
    generated_content: str
    sources: List[Dict]
    error: Optional[str] = None

# Mapping from form field keys to section labels used in prompts
FIELD_TO_SECTION = {
    "projectTitle": "Project Title",
    # NSF fields
    "researchSpaceFacilities": "Research Space and Facilities",
    "coreInstrumentation": "Core Instrumentation",
    "computingDataResources": "Computing and Data Resources",
    "internalFacilitiesNYU": "Internal Facilities (NYU)",
    "externalFacilitiesOther": "External Facilities (Other Institutions)",
    "specialInfrastructure": "Special Infrastructure",
    # NIH fields
    "laboratory": "Laboratory",
    "animal": "Animal",
    "computer": "Computer",
    "office": "Office",
    "clinical": "Clinical",
    "other": "Other",
    "equipment": "Equipment"
}

def get_section_labels(sponsor: str) -> List[str]:
    """Get section labels based on sponsor type"""
    if sponsor == "NSF":
        # NSF has 4 consolidated sections
        return [
            "1. Project Title",
            "2. Facilities",
            "3. Major Equipment",
            "4. Other Resources"
        ]
    else:  # NIH
        # NIH keeps individual sections
        return [
            "1. Project Title",
            "2. Laboratory",
            "3. Animal",
            "4. Computer",
            "5. Office",
            "6. Clinical",
            "7. Other",
            "8. Equipment"
        ]

def facilities_web_search(query: str, request: Request, user) -> tuple[str, List[str], List[float]]:
    """
     web search for facilities - uses existing NAGA Tavily configuration
    """
    try:
        # Check if web search is enabled for this user
        if not request.app.state.config.ENABLE_RAG_WEB_SEARCH.get(user.email):
            logging.info(f"Web search is disabled for user: {user.email}")
            return "", [], []
        
        if not request.app.state.config.TAVILY_API_KEY:
            logging.warning("TAVILY_API_KEY not configured in NAGA")
            return "", [], []
        
        tavily_api_key = request.app.state.config.TAVILY_API_KEY.get(user.email)
        if not tavily_api_key:
            logging.warning(f"No Tavily API key found for user: {user.email} (checked user config and group admin config)")
            return "", [], []
        
        logging.info(f"Starting facilities web search for query: {query}")
        
        allowed_domains = request.app.state.config.RAG_WEB_SEARCH_DOMAIN_FILTER_LIST.get(user.email)
        blocklist = request.app.state.config.RAG_WEB_SEARCH_WEBSITE_BLOCKLIST.get(user.email) or []
        
        snippets = []
        urls = []
        scores = []
        
        if allowed_domains:
            # Search specific domains
            for domain in allowed_domains:
                logging.info(f"Searching domain: {domain}")
                try:
                    site_query = f"site:{domain} {query}"
                    results = search_tavily(
                        api_key=tavily_api_key,
                        query=site_query,
                        count=3,
                        filter_list=None,
                        blocklist=blocklist
                    )
                    logging.info(f"Search results for {domain}: {len(results) if results else 0} results")
                    
                    if results:
                        for result in results:
                            url = result.link.strip()
                            content = result.snippet.strip()
                            score = result.score if hasattr(result, 'score') else 0.1
                            logging.info(f"Found result: {url[:50]}... with score: {score}")
                            
                            # Parse domain and validate it's from allowed domains
                            from urllib.parse import urlparse
                            parsed_domain = urlparse(url).netloc
                            
                            # More strict domain validation - must end with allowed domains
                            is_allowed_domain = any(parsed_domain.endswith(allowed) for allowed in allowed_domains)
                            is_blocked = any(url.startswith(bad) for bad in blocklist)
                            
                            if is_allowed_domain and not is_blocked:
                                snippets.append(content)
                                urls.append(url)
                                scores.append(score)
                                logging.info(f"Added to results: {url} with score: {score} (domain: {parsed_domain})")
                            else:
                                logging.info(f"Blocked by filter: {url} (domain: {parsed_domain}, allowed: {is_allowed_domain}, blocked: {is_blocked})")
                    else:
                        logging.info(f"No results for domain: {domain}")
                except Exception as e:
                    logging.error(f"Error searching domain {domain}: {e}")
        else:
            # No domain filter - search the entire web (like regular chat)
            logging.info(f"No domain filter configured, searching entire web for query: {query}")
            try:
                results = search_tavily(
                    api_key=tavily_api_key,
                    query=query,
                    count=3,
                    filter_list=None,
                    blocklist=blocklist
                )
                logging.info(f"Web search results: {len(results) if results else 0} results")
                
                if results:
                    for result in results:
                        url = result.link.strip()
                        content = result.snippet.strip()
                        score = result.score if hasattr(result, 'score') else 0.1
                        logging.info(f"Found result: {url[:50]}... with score: {score}")
                        
                        # Parse domain and validate it's not blocked
                        from urllib.parse import urlparse
                        parsed_domain = urlparse(url).netloc
                        is_blocked = any(url.startswith(bad) for bad in blocklist)
                        
                        if not is_blocked:
                            snippets.append(content)
                            urls.append(url)
                            scores.append(score)
                            logging.info(f"Added to results: {url} with score: {score}")
                        else:
                            logging.info(f"Blocked by blocklist: {url}")
                else:
                    logging.info(f"No web search results found")
            except Exception as e:
                logging.error(f"Error in web search: {e}")
        
        return "\n\n".join(snippets), urls, scores
        
    except Exception as e:
        logging.error(f"Facilities web search failed: {e}")
        return f"Web search failed: {e}", [], []

def facilities_web_search_specific_sites(query: str, allowed_sites: List[str], request: Request, user) -> tuple[str, List[str], List[float]]:
    """
    web search within specific sites - uses existing NAGA Tavily configuration
    """
    try:
       
        if not request.app.state.config.TAVILY_API_KEY:
            logging.warning("TAVILY_API_KEY not configured in NAGA")
            return "", [], []
        
        
        tavily_api_key = request.app.state.config.TAVILY_API_KEY.get(user.email)
        if not tavily_api_key:
            logging.warning(f"No Tavily API key found for user: {user.email} (checked user config and group admin config)")
            return "", [], []
        
        logging.info(f"Starting specific sites web search for query: {query}, sites: {allowed_sites}")
        
        blocklist = request.app.state.config.RAG_WEB_SEARCH_WEBSITE_BLOCKLIST.get(user.email) or []
        
        snippets = []
        urls = []
        scores = []
        
        
        for site in allowed_sites:
            
            # Use the exact URL for internal facilities
            site_query = f"site:{site} {query}"
            logging.info(f"Searching site: {site} with query: {site_query}")
            results = search_tavily(
                api_key=tavily_api_key,
                query=site_query,
                count=3,
                filter_list=None,
                blocklist=blocklist
            )
            logging.info(f"Found {len(results) if results else 0} results for site: {site}")
            
            for result in results:
                url = result.link
                content = result.snippet.strip()
                score = result.score if hasattr(result, 'score') else 0.1
                
                # Check if the result URL starts with any of the allowed site paths
                # This ensures ONLY pages from the specific configured URLs are accepted
                is_from_allowed_site = False
                for allowed_site in allowed_sites:
                    # Normalize URLs for comparison (remove trailing slashes)
                    normalized_allowed = allowed_site.rstrip('/')
                    normalized_result = url.rstrip('/')
                    
                    # Check if result URL starts with the allowed site path
                    if normalized_result.startswith(normalized_allowed) or normalized_result == normalized_allowed:
                        is_from_allowed_site = True
                        logging.info(f"Result URL {url[:80]}... matches allowed site {allowed_site}")
                        break
                
                if is_from_allowed_site:
                    snippets.append(content)
                    urls.append(url)
                    scores.append(score)
                    logging.info(f"Added result from allowed site: {url[:100]}...")
        
        return "\n\n".join(snippets), urls, scores
        
    except Exception as e:
        logging.error(f"Facilities specific site web search failed: {e}")
        return f"Web search failed: {e}", [], []

def search_knowledge_base(query: str, user_id: str, request: Request, model, k: int = 10) -> List[tuple]:
    """Model-dependent knowledge base search"""
    try:
        # Get model knowledge configuration (same as regular chat)
        model_knowledge = model.get("info", {}).get("meta", {}).get("knowledge", False)
        
        if not model_knowledge:
            logging.warning("No knowledge bases configured for this model")
            return []
        
        # Extract collection names from model configuration
        collection_names = []
        for item in model_knowledge:
            if item.get("collection_name"):
                collection_names.append(item.get("collection_name"))
            elif item.get("collection_names"):
                collection_names.extend(item.get("collection_names", []))
            elif item.get("id"):
                collection_names.append(item.get("id"))
        
        if not collection_names:
            logging.warning("No knowledge base collections found in model configuration")
            return []
        
        try:
            embedding_function = request.app.state.EMBEDDING_FUNCTION
            query_embedding = embedding_function(query, user_id)
        except Exception as e:
            logging.error(f"Failed to generate embeddings: {e}")
            return []
        
        
        results = []
        for collection_name in collection_names:
            try:
               
                search_results = VECTOR_DB_CLIENT.search(
                    collection_name=collection_name,
                    vectors=[query_embedding],
                    limit=k * 3  
                )
                
                if search_results and hasattr(search_results, 'documents') and search_results.documents:
                    documents = search_results.documents[0]
                    metadatas = search_results.metadatas[0] if search_results.metadatas else []
                    distances = search_results.distances[0] if hasattr(search_results, 'distances') and search_results.distances else []
                    
                    for i, doc in enumerate(documents):
                        if i < len(metadatas) and metadatas[i]:
                            
                            metadata = metadatas[i]
                            source = 'unknown'
                            
                            logging.info(f"Metadata for document {i}: {metadata}")
                            
                            
                            for key in ['source', 'name', 'filename', 'file_name']:
                                if key in metadata:
                                    source = metadata[key]
                                    logging.info(f"Found source '{source}' using key '{key}'")
                                    break
                        else:
                            source = 'unknown'
                            logging.warning(f"No metadata found for document {i}")
                        
                        # Get real distance/score if available
                        real_score = distances[i] if i < len(distances) else 0.15
                        results.append((doc, source, real_score))
                        
            except Exception as e:
                logging.warning(f"Failed to search collection {collection_name}: {e}")
                continue
        
        return results[:k]
        
    except Exception as e:
        logging.error(f"Knowledge base search failed: {e}")
        return []

async def call_llm_direct(prompt: str, model: str, user, request: Request) -> str:
    """Direct LLM call using existing chat completion system"""
    try:
        from open_webui.utils.chat import generate_chat_completion
        
        form_data = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False
        }
        
        completion_response = await generate_chat_completion(
            request=request,
            form_data=form_data,
            user=user,
            bypass_filter=True
        )
        

        if isinstance(completion_response, dict) and "choices" in completion_response:
            return completion_response["choices"][0]["message"]["content"].strip()
        else:
            raise Exception(f"Unexpected LLM response format: {type(completion_response)}")
            
    except Exception as e:
        logging.error(f"Direct LLM call failed: {e}")
        raise Exception(f"LLM call failed: {str(e)}")

@router.post("/generate", response_model=FacilitiesResponse)
async def generate_facilities_response(request: Request, form_data: FacilitiesRequest, user=Depends(get_verified_user)):
    """
    Generate facilities response
    Only works when facilities is enabled for the user
    """
    if not request.app.state.config.ENABLE_FACILITIES.get(user.email):
        raise HTTPException(status_code=403, detail="Facilities feature is not enabled for this user")
    try:
        logging.info(f"DEBUG: Received facilities request with files: {form_data.files}")
        logging.info(f"DEBUG: Files count: {len(form_data.files) if form_data.files else 0}")
        if form_data.files:
            for i, file in enumerate(form_data.files):
                logging.info(f"DEBUG: File {i}: {file}")
        if form_data.sponsor not in ["NSF", "NIH"]:
            raise HTTPException(status_code=400, detail="Invalid sponsor. Must be 'NSF' or 'NIH'")
        
        # Get model information for knowledge base access
        model = request.app.state.MODELS.get(form_data.model)
        if not model:
            raise HTTPException(status_code=400, detail="Model not found")
        
        section_labels = get_section_labels(form_data.sponsor)
        
        # Mapping from form fields to section labels (1-to-1 for both NSF and NIH)
        if form_data.sponsor == "NSF":
            # NSF: 1-to-1 mapping to individual sections (will be grouped in response)
            field_mappings = {
                "projectTitle": "Project Title",
                "researchSpaceFacilities": "Research Space and Facilities", 
                "coreInstrumentation": "Core Instrumentation",
                "computingDataResources": "Computing and Data Resources",
                "internalFacilitiesNYU": "Internal Facilities (NYU)",
                "externalFacilitiesOther": "External Facilities (Other Institutions)",
                "specialInfrastructure": "Special Infrastructure"
            }
        else:
            # NIH: Individual sections
            field_mappings = {
                "projectTitle": "1. Project Title",
                "laboratory": "2. Laboratory",
                "animal": "3. Animal",
                "computer": "4. Computer",
                "office": "5. Office",
                "clinical": "6. Clinical",
                "other": "7. Other",
                "equipment": "8. Equipment"
            }
        
        user_inputs = {}
        for form_key, section_label in field_mappings.items():
            if form_key in form_data.form_data:
                text = form_data.form_data[form_key].strip()
                if text:
                    user_inputs[section_label] = text
        
        # For NSF: Create mapping from individual sections to NSF response headers
        nsf_header_mapping = {}
        if form_data.sponsor == "NSF":
            nsf_header_mapping = {
                "Project Title": "1. Project Title",
                "Research Space and Facilities": "2. Facilities",
                "Core Instrumentation": "3. Major Equipment",
                "Computing and Data Resources": "3. Major Equipment",
                "Internal Facilities (NYU)": "4. Other Resources",
                "External Facilities (Other Institutions)": "4. Other Resources",
                "Special Infrastructure": "4. Other Resources"
            }
        
        if not any(user_inputs.values()):
            return FacilitiesResponse(
                success=False,
                message="No input provided in form fields",
                sections={},
                sources=[],
                error="Please fill in at least one form field"
            )
        
        # Import the file processing function from middleware (moved outside loop for efficiency)
        from open_webui.utils.middleware import chat_completion_files_handler
        
        section_outputs = {}
        all_sources = []
        cited_sources = set()  # Track which sources are actually cited in the text
        # Using source names directly - frontend will handle the mapping  
        
        for section, user_text in user_inputs.items():
            if not user_text:
                continue
            
            logging.info(f"Processing section: {section}")
            logging.info(f"Web search enabled flag: {form_data.web_search_enabled}")
            
            # Improved query construction: use user text directly for better semantic matching
            # The section context is preserved in the prompt later, so we don't need it in the query
            query = user_text
            
            # (query expansion, hybrid search, reranking, TOP_K config)
            all_file_results = []
            
            # Get model knowledge configuration
            model_knowledge = model.get("info", {}).get("meta", {}).get("knowledge", False)
            
            # Combine knowledge base collections and uploaded files
            files_to_process = []
            
            # Add knowledge base collections as "files"
            if model_knowledge:
                knowledge_files = []
                for item in model_knowledge:
                    if item.get("collection_name"):
                        knowledge_files.append({
                            "id": item.get("collection_name"),
                            "name": item.get("name"),
                            "legacy": True,
                        })
                    elif item.get("collection_names"):
                        knowledge_files.append({
                            "name": item.get("name"),
                            "type": "collection",
                            "collection_names": item.get("collection_names"),
                            "legacy": True,
                        })
                    else:
                        knowledge_files.append(item)
                
                files_to_process.extend(knowledge_files)
                logging.info(f"Added {len(knowledge_files)} knowledge base collections for section '{section}'")
            
            # Add uploaded files
            if form_data.files:
                files_to_process.extend(form_data.files)
                logging.info(f"Added {len(form_data.files)} uploaded files for section '{section}'")
            
            # Process all files (knowledge base + uploaded) together with advanced pipeline
            if files_to_process:
                try:
                    # Create section-specific body for file processing
                    section_specific_body = {
                        "model": form_data.model,
                        "metadata": {
                            "files": files_to_process  # Knowledge base + uploaded files together
                        },
                        "messages": [{"role": "user", "content": user_text}]  # Section-specific query
                    }
                    
                    processed_body, file_flags = await chat_completion_files_handler(request, section_specific_body, user)
                    all_file_sources = file_flags.get("sources", [])
                    
                    # Extract documents with actual distance scores
                    for source in all_file_sources:
                        documents = source.get("document", [])
                        distances = source.get("distances", [])
                        metadata = source.get("metadata", [])
                        
                        # Extract source name from metadata (filename) first, fallback to file object name
                        source_name = "unknown"
                        if metadata and isinstance(metadata, list) and len(metadata) > 0:
                            # Get first metadata entry (should have filename/source info)
                            first_metadata = metadata[0] if isinstance(metadata[0], dict) else {}
                            # Try metadata fields that contain actual filenames
                            for key in ['source', 'name', 'filename', 'file_name']:
                                if key in first_metadata and first_metadata[key]:
                                    source_name = first_metadata[key]
                                    break
                        
                        # Fallback to file object name if metadata doesn't have filename
                        if source_name == "unknown":
                            file_obj = source.get("source", {})
                            if isinstance(file_obj, dict):
                                source_name = file_obj.get("name", file_obj.get("id", "unknown"))
                        
                        # Use actual distance scores if available, otherwise use a default
                        for i, doc in enumerate(documents):
                            distance = distances[i] if i < len(distances) and distances[i] is not None else 0.15
                            all_file_results.append((doc, source_name, distance))
                    
                    logging.info(f"Processed {len(all_file_sources)} sources (KB + uploaded) for section '{section}'")
                    logging.info(f"Added {len(all_file_results)} total chunks for {section}")
                    
                except Exception as e:
                    logging.error(f"Error processing files for section '{section}': {e}")
                    # Continue without files for this section if processing fails
            
            # All results are now in all_file_results (knowledge base + uploaded files combined)
            all_search_results = all_file_results
            
             
            retrieved_chunks = []
            section_sources = []
            
            
            source_groups = {}
            for doc, source, score in all_search_results:
                if source not in source_groups:
                    source_groups[source] = []
                source_groups[source].append((doc, score))
            
            for source, docs_with_scores in source_groups.items():
                
                logging.info(f"Processing source '{source}' with {len(docs_with_scores)} documents, using source name directly")
                
                docs = [doc for doc, score in docs_with_scores]
                real_distances = [score for doc, score in docs_with_scores]
                
                for doc in docs:
                    retrieved_chunks.append(f"<source><source_id>{source}</source_id><source_context>{doc}</source_context></source>")
                
                section_sources.append({
                    "source": {"id": source, "name": source},
                    "document": docs,  
                    "metadata": [{"source": source, "name": source}] * len(docs),
                    "distances": real_distances 
                })
                # No need to increment since we're using source names directly
            
            retrieved_texts_with_sources = "\n\n".join(retrieved_chunks) if retrieved_chunks else "No relevant PDF documents found in knowledge base."
            all_sources.extend(section_sources)
            
            logging.info(f"Added {len(section_sources)} sources to all_sources for {section}")
            logging.info(f"Sources added: {[s.get('source', {}).get('name', 'unknown') for s in section_sources]}")
            
            logging.info(f"Found {len(all_search_results)} total chunks for {section} (KB + uploaded files processed together with advanced pipeline)")
            logging.info(f"All sources found: {[source for _, source, _ in all_search_results]}")
            
            web_content = ""
            web_links = []
            web_scores = []
            web_source_tags = []
            
            if form_data.web_search_enabled:
                logging.info(f"Web search enabled for section: {section}")
                # Check if this is Internal Facilities (works for both NSF and NIH)
                is_internal_facilities = (
                    section == "5a. Internal Facilities (NYU)" or 
                    section == "Internal Facilities (NYU)"
                )
                
                if is_internal_facilities:
                    allowed_sites = request.app.state.config.RAG_WEB_SEARCH_INTERNAL_FACILITIES_SITES.get(user.email)
                    logging.info(f"Using specific sites search for NYU Internal Facilities")
                    logging.info(f"Allowed sites: {allowed_sites}")
                    
                    if allowed_sites:
                        web_content, web_links, web_scores = facilities_web_search_specific_sites(
                            query, allowed_sites, request, user
                        )
                    else:
                        logging.warning("No allowed sites configured for NYU Internal Facilities, falling back to standard search")
                        web_content, web_links, web_scores = facilities_web_search(query, request, user)
                else:
                    # Use improved query (user_text) for standard web search
                    logging.info(f"Using standard web search for section: {section}")
                    web_content, web_links, web_scores = facilities_web_search(query, request, user)
                
                logging.info(f"Web search results - Content length: {len(web_content) if web_content else 0}, Links: {len(web_links) if web_links else 0}")
                
                if web_content and web_links:
                    
                    web_content_parts = web_content.split('\n\n')
                    
                    for i, (content_part, url, score) in enumerate(zip(web_content_parts, web_links, web_scores)):
                        if content_part.strip() and url.strip():
                            # Use current source_index for this web source
                            logging.info(f"Processing web source '{url}' using URL directly with score: {score}")
                            web_source_tags.append(f"<source><source_id>{url}</source_id><source_context>{content_part.strip()}</source_context></source>")
                            

                            web_distances = [score]  # Use real score from web search
                            
                            # Create separate source for each URL (like existing NAGA system)
                            all_sources.append({
                                "source": {"id": url, "name": url, "url": url},  # Use URL as both name and url
                                "document": [content_part.strip()],
                                "metadata": [{"source": url, "name": url}],
                                "distances": web_distances  # Use real distances/scores
                            })
                            logging.info(f"Added web source to all_sources: {url}")
                            # No need to increment since we're using source names directly
                
                logging.info(f"Found {len(web_links)} web sources for {section}")
                logging.info(f"Web URLs found: {web_links}")
            else:
                logging.info(f"Web search disabled for {section}")
            
            pdf_sources = "\n\n".join(retrieved_chunks) if retrieved_chunks else "No relevant PDF documents found in knowledge base."
            # Only include web sources if web search is enabled
            if form_data.web_search_enabled:
                web_sources = "\n\n".join(web_source_tags) if web_source_tags else "No relevant web sources found."
            else:
                web_sources = "Web search is disabled for this request."
            
            # For NSF: Individual sections are processed separately
            # The grouping under NSF headers happens in the response formatting
            section_specific_instructions = ""
            integration_instructions = ""
            
            # Use the section name as-is for the prompt (will be grouped later in response for NSF)
            prompt = FACILITIES_PROMPT.format(
                sponsor=form_data.sponsor,
                section=section,
                section_specific_instructions=section_specific_instructions,
                user_input=user_text,
                retrieved_chunks=pdf_sources,
                web_snippets=web_sources,
                integration_instructions=integration_instructions
            )
            
            try:
                generated_content = await call_llm_direct(prompt, form_data.model, user, request)
                
                import re
                cleaned_content = generated_content.strip()
                
                if cleaned_content.startswith('Error:') or 'Connection aborted' in cleaned_content:
                    logging.error(f"LLM returned error for {section}: {cleaned_content}")
                    raise Exception(f"LLM generation failed for {section}: {cleaned_content}")
                else:
                    section_outputs[section] = cleaned_content
                    
                    # Extract cited sources from the generated text
                    import re
                    # Find all citations in the format [source_name] or [url]
                    citations = re.findall(r'\[([^\]]+)\]', cleaned_content)
                    for citation in citations:
                        # Only add if it looks like a source (contains .pdf, .doc, http, or is a filename)
                        # But exclude web URLs if web search is disabled
                        if ('.pdf' in citation or '.doc' in citation or 
                            (citation.startswith('http') and form_data.web_search_enabled) or 
                            (len(citation) > 10 and not citation.startswith('http'))):  # Likely a filename, not a URL
                            cited_sources.add(citation)
                    
                    logging.info(f"Generated content for {section}: {len(cleaned_content)} chars")
                    logging.info(f"All citations found in {section}: {citations}")
                    # Filter citations based on web search status
                    if form_data.web_search_enabled:
                        valid_citations = [c for c in citations if '.pdf' in c or '.doc' in c or c.startswith('http') or len(c) > 10]
                    else:
                        valid_citations = [c for c in citations if '.pdf' in c or '.doc' in c or (len(c) > 10 and not c.startswith('http'))]
                    logging.info(f"Valid source citations in {section}: {valid_citations}")
                    
                    # Separate PDF and web citations for better debugging
                    pdf_citations = [c for c in valid_citations if '.pdf' in c or '.doc' in c]
                    web_citations = [c for c in valid_citations if c.startswith('http')]
                    logging.info(f"PDF citations in {section}: {pdf_citations}")
                    logging.info(f"Web citations in {section}: {web_citations}")
                    logging.info(f"Content preview: {cleaned_content[:200]}...")
                
            except Exception as e:
                logging.error(f"Failed to generate content for {section}: {e}")
                raise Exception(f"Failed to generate content for {section}: {str(e)}")
        
        if not section_outputs:
            return FacilitiesResponse(
                success=False,
                message="No sections generated",
                content="",
                sections={},
                sources=[],
                error="Failed to generate any sections"
            )
        
        # For NSF: Group individual sections under NSF headers
        if form_data.sponsor == "NSF":
            # Group sections by NSF header
            grouped_sections = {}
            for individual_section, content in section_outputs.items():
                nsf_header = nsf_header_mapping.get(individual_section)
                if nsf_header:
                    if nsf_header not in grouped_sections:
                        grouped_sections[nsf_header] = []
                    grouped_sections[nsf_header].append({
                        "subsection": individual_section,
                        "content": content
                    })
                else:
                    logging.warning(f"No NSF header mapping found for section: {individual_section}")
            
            # Build formatted response with NSF headers and subsections
            formatted_sections = {}
            content = f"# Facilities Response for {form_data.sponsor}\n\n"
            
            # Define NSF header order and subsection order within each header
            nsf_header_order = [
                "1. Project Title",
                "2. Facilities",
                "3. Major Equipment",
                "4. Other Resources"
            ]
            
            # Define which headers should always use subsection headings (even for single subsections)
            headers_with_subsections = ["3. Major Equipment", "4. Other Resources"]
            
            # Define subsection order for headers with multiple subsections
            subsection_order = {
                "3. Major Equipment": ["Core Instrumentation", "Computing and Data Resources"],
                "4. Other Resources": ["Internal Facilities (NYU)", "External Facilities (Other Institutions)", "Special Infrastructure"]
            }
            
            for nsf_header in nsf_header_order:
                # Only include headers that have at least one subsection filled
                # If no subsections are filled for a header, it will be skipped entirely
                if nsf_header in grouped_sections:
                    subsections = grouped_sections[nsf_header]
                    
                    if nsf_header in headers_with_subsections:
                        # Sort subsections according to defined order
                        if nsf_header in subsection_order:
                            ordered_subsections = []
                            subsection_dict = {sub['subsection']: sub for sub in subsections}
                            for ordered_sub_name in subsection_order[nsf_header]:
                                if ordered_sub_name in subsection_dict:
                                    ordered_subsections.append(subsection_dict[ordered_sub_name])
                            # Add any remaining subsections not in the order list
                            for sub in subsections:
                                if sub['subsection'] not in subsection_order[nsf_header]:
                                    ordered_subsections.append(sub)
                            subsections = ordered_subsections
                        
                        formatted_content = ""
                        for sub in subsections:
                            formatted_content += f"### {sub['subsection']}\n\n{sub['content']}\n\n"
                        formatted_sections[nsf_header] = formatted_content.strip()
                        content += f"## {nsf_header}\n\n{formatted_content}\n\n"
                    else:
                        # For single subsection headers (1. Project Title, 2. Facilities), include content directly (no ### heading)
                        formatted_sections[nsf_header] = subsections[0]['content']
                        content += f"## {nsf_header}\n\n{subsections[0]['content']}\n\n"
            
            # Update section_outputs to use NSF headers for response
            section_outputs = formatted_sections
        else:
            # NIH: Keep individual sections as-is
            content = f"# Facilities Response for {form_data.sponsor}\n\n"
            for section_name, section_content in section_outputs.items():
                content += f"## {section_name}\n\n{section_content}\n\n"
        
        # Filter sources to only include those that were actually cited in the text
        filtered_sources = []
        for source in all_sources:
            source_name = source.get('source', {}).get('name', '')
            if source_name in cited_sources:
                filtered_sources.append(source)
        
        formatted_sources = filtered_sources
        logging.info(f"Total sources found: {len(all_sources)}")
        logging.info(f"Sources actually cited: {len(cited_sources)}")
        logging.info(f"Final sources count: {len(formatted_sources)}")
        logging.info(f"Cited sources: {list(cited_sources)}")
        logging.info(f"Final sources: {[s.get('source', {}).get('name', 'unknown') for s in formatted_sources]}")
        
        # Facilities response is added to current chat by frontend addFacilitiesResponseToChat function
        
        return FacilitiesResponse(
            success=True,
            message=f"Successfully generated {len(section_outputs)} sections for {form_data.sponsor}",
            content=content,
            sections=section_outputs, 
            sources=formatted_sources,
            error=None  # No error for successful generation
        )
        
    except Exception as e:
        logging.error(f"Error in facilities generation: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@router.post("/generate-section", response_model=SingleSectionResponse)
async def generate_facilities_section(request: Request, form_data: SingleSectionRequest, user=Depends(get_verified_user)):
    """
    Generate facilities content for a single section.
    Called once per section so the frontend can display results progressively.
    """
    if not request.app.state.config.ENABLE_FACILITIES.get(user.email):
        raise HTTPException(status_code=403, detail="Facilities feature is not enabled for this user")

    try:
        if form_data.sponsor not in ["NSF", "NIH"]:
            raise HTTPException(status_code=400, detail="Invalid sponsor. Must be 'NSF' or 'NIH'")

        section = FIELD_TO_SECTION.get(form_data.section_key)
        if not section:
            raise HTTPException(status_code=400, detail=f"Invalid section key: {form_data.section_key}")

        user_text = form_data.section_text.strip()
        if not user_text:
            return SingleSectionResponse(
                success=False,
                section_key=form_data.section_key,
                section_label=section,
                generated_content="",
                sources=[],
                error="Empty section text"
            )

        model = request.app.state.MODELS.get(form_data.model)
        if not model:
            raise HTTPException(status_code=400, detail="Model not found")

        from open_webui.utils.middleware import chat_completion_files_handler
        import re

        logging.info(f"Processing section: {section}")
        logging.info(f"Web search enabled flag: {form_data.web_search_enabled}")
        query = user_text
        all_file_results = []
        all_sources = []
        cited_sources = set()

        # --- File processing (knowledge base + uploaded files) ---
        model_knowledge = model.get("info", {}).get("meta", {}).get("knowledge", False)
        files_to_process = []

        if model_knowledge:
            knowledge_files = []
            for item in model_knowledge:
                if item.get("collection_name"):
                    knowledge_files.append({
                        "id": item.get("collection_name"),
                        "name": item.get("name"),
                        "legacy": True,
                    })
                elif item.get("collection_names"):
                    knowledge_files.append({
                        "name": item.get("name"),
                        "type": "collection",
                        "collection_names": item.get("collection_names"),
                        "legacy": True,
                    })
                else:
                    knowledge_files.append(item)
            files_to_process.extend(knowledge_files)
            logging.info(f"Added {len(knowledge_files)} knowledge base collections for section '{section}'")

        if form_data.files:
            files_to_process.extend(form_data.files)
            logging.info(f"Added {len(form_data.files)} uploaded files for section '{section}'")

        if files_to_process:
            try:
                section_specific_body = {
                    "model": form_data.model,
                    "metadata": {"files": files_to_process},
                    "messages": [{"role": "user", "content": user_text}]
                }
                processed_body, file_flags = await chat_completion_files_handler(request, section_specific_body, user)
                all_file_sources = file_flags.get("sources", [])

                for source in all_file_sources:
                    documents = source.get("document", [])
                    distances = source.get("distances", [])
                    metadata = source.get("metadata", [])

                    source_name = "unknown"
                    if metadata and isinstance(metadata, list) and len(metadata) > 0:
                        first_metadata = metadata[0] if isinstance(metadata[0], dict) else {}
                        for key in ['source', 'name', 'filename', 'file_name']:
                            if key in first_metadata and first_metadata[key]:
                                source_name = first_metadata[key]
                                break

                    if source_name == "unknown":
                        file_obj = source.get("source", {})
                        if isinstance(file_obj, dict):
                            source_name = file_obj.get("name", file_obj.get("id", "unknown"))

                    for i, doc in enumerate(documents):
                        distance = distances[i] if i < len(distances) and distances[i] is not None else 0.15
                        all_file_results.append((doc, source_name, distance))

                logging.info(f"Processed {len(all_file_sources)} sources (KB + uploaded) for section '{section}'")
                logging.info(f"Added {len(all_file_results)} total chunks for {section}")
            except Exception as e:
                logging.error(f"Error processing files for section '{section}': {e}")

        all_search_results = all_file_results

        # --- Source grouping ---
        retrieved_chunks = []
        section_sources = []

        source_groups = {}
        for doc, source, score in all_search_results:
            if source not in source_groups:
                source_groups[source] = []
            source_groups[source].append((doc, score))

        for source, docs_with_scores in source_groups.items():
            logging.info(f"Processing source '{source}' with {len(docs_with_scores)} documents, using source name directly")
            
            docs = [doc for doc, score in docs_with_scores]
            real_distances = [score for doc, score in docs_with_scores]

            for doc in docs:
                retrieved_chunks.append(f"<source><source_id>{source}</source_id><source_context>{doc}</source_context></source>")

            section_sources.append({
                "source": {"id": source, "name": source},
                "document": docs,
                "metadata": [{"source": source, "name": source}] * len(docs),
                "distances": real_distances
            })

        all_sources.extend(section_sources)
        
        logging.info(f"Added {len(section_sources)} sources to all_sources for {section}")
        logging.info(f"Sources added: {[s.get('source', {}).get('name', 'unknown') for s in section_sources]}")
        logging.info(f"Found {len(all_search_results)} total chunks for {section} (KB + uploaded files processed together with advanced pipeline)")
        logging.info(f"All sources found: {[source for _, source, _ in all_search_results]}")

        # --- Web search ---
        web_content = ""
        web_links = []
        web_scores = []
        web_source_tags = []
        
        if form_data.web_search_enabled:
            logging.info(f"Web search enabled for section: {section}")
            is_internal_facilities = (
                section == "5a. Internal Facilities (NYU)" or
                section == "Internal Facilities (NYU)"
            )

            if is_internal_facilities:
                allowed_sites = request.app.state.config.RAG_WEB_SEARCH_INTERNAL_FACILITIES_SITES.get(user.email)
                logging.info(f"Using specific sites search for NYU Internal Facilities")
                logging.info(f"Allowed sites: {allowed_sites}")
                
                if allowed_sites:
                    web_content, web_links, web_scores = facilities_web_search_specific_sites(
                        query, allowed_sites, request, user
                    )
                else:
                    logging.warning("No allowed sites configured for NYU Internal Facilities, falling back to standard search")
                    web_content, web_links, web_scores = facilities_web_search(query, request, user)
            else:
                logging.info(f"Using standard web search for section: {section}")
                web_content, web_links, web_scores = facilities_web_search(query, request, user)
            
            logging.info(f"Web search results - Content length: {len(web_content) if web_content else 0}, Links: {len(web_links) if web_links else 0}")

            if web_content and web_links:
                web_content_parts = web_content.split('\n\n')
                for i, (content_part, url, score) in enumerate(zip(web_content_parts, web_links, web_scores)):
                    if content_part.strip() and url.strip():
                        logging.info(f"Processing web source '{url}' using URL directly with score: {score}")
                        web_source_tags.append(
                            f"<source><source_id>{url}</source_id><source_context>{content_part.strip()}</source_context></source>"
                        )
                        all_sources.append({
                            "source": {"id": url, "name": url, "url": url},
                            "document": [content_part.strip()],
                            "metadata": [{"source": url, "name": url}],
                            "distances": [score]
                        })
                        logging.info(f"Added web source to all_sources: {url}")
            
            logging.info(f"Found {len(web_links)} web sources for {section}")
            logging.info(f"Web URLs found: {web_links}")
        else:
            logging.info(f"Web search disabled for {section}")

        # --- Build prompt ---
        pdf_sources = "\n\n".join(retrieved_chunks) if retrieved_chunks else "No relevant PDF documents found in knowledge base."
        if form_data.web_search_enabled:
            web_sources = "\n\n".join(web_source_tags) if web_source_tags else "No relevant web sources found."
        else:
            web_sources = "Web search is disabled for this request."

        prompt = FACILITIES_PROMPT.format(
            sponsor=form_data.sponsor,
            section=section,
            section_specific_instructions="",
            user_input=user_text,
            retrieved_chunks=pdf_sources,
            web_snippets=web_sources,
            integration_instructions=""
        )

        # --- LLM call ---
        generated_content = await call_llm_direct(prompt, form_data.model, user, request)
        cleaned_content = generated_content.strip()

        if cleaned_content.startswith('Error:') or 'Connection aborted' in cleaned_content:
            raise Exception(f"LLM generation failed for {section}: {cleaned_content}")

        # --- Citation extraction & source filtering ---
        citations = re.findall(r'\[([^\]]+)\]', cleaned_content)
        for citation in citations:
            if ('.pdf' in citation or '.doc' in citation or 
                (citation.startswith('http') and form_data.web_search_enabled) or 
                (len(citation) > 10 and not citation.startswith('http'))):
                cited_sources.add(citation)
        
        logging.info(f"Generated content for {section}: {len(cleaned_content)} chars")
        logging.info(f"All citations found in {section}: {citations}")
        # Filter citations based on web search status
        if form_data.web_search_enabled:
            valid_citations = [c for c in citations if '.pdf' in c or '.doc' in c or c.startswith('http') or len(c) > 10]
        else:
            valid_citations = [c for c in citations if '.pdf' in c or '.doc' in c or (len(c) > 10 and not c.startswith('http'))]
        logging.info(f"Valid source citations in {section}: {valid_citations}")
        
        # Separate PDF and web citations for better debugging
        pdf_citations = [c for c in valid_citations if '.pdf' in c or '.doc' in c]
        web_citations = [c for c in valid_citations if c.startswith('http')]
        logging.info(f"PDF citations in {section}: {pdf_citations}")
        logging.info(f"Web citations in {section}: {web_citations}")
        logging.info(f"Content preview: {cleaned_content[:200]}...")

        filtered_sources = []
        for source in all_sources:
            source_name = source.get('source', {}).get('name', '')
            if source_name in cited_sources:
                filtered_sources.append(source)

        logging.info(f"Total sources found: {len(all_sources)}")
        logging.info(f"Sources actually cited: {len(cited_sources)}")
        logging.info(f"Final sources count: {len(filtered_sources)}")

        return SingleSectionResponse(
            success=True,
            section_key=form_data.section_key,
            section_label=section,
            generated_content=cleaned_content,
            sources=filtered_sources,
            error=None
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"[generate-section] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@router.get("/sections/{sponsor}")
async def get_facilities_sections(sponsor: str, request: Request, user=Depends(get_verified_user)):
    """
    Get the section labels for a specific sponsor
    Only works when facilities is enabled for the user
    """
    if not request.app.state.config.ENABLE_FACILITIES.get(user.email):
        raise HTTPException(status_code=403, detail="Facilities feature is not enabled for this user")
    try:
        if sponsor not in ["NSF", "NIH"]:
            raise HTTPException(status_code=400, detail="Invalid sponsor. Must be 'NSF' or 'NIH'")
        
        section_labels = get_section_labels(sponsor)
        
        return {
            "success": True,
            "sponsor": sponsor,
            "sections": section_labels
        }
        
    except Exception as e:
        logging.error(f"Error getting sections for {sponsor}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get sections: {str(e)}")


class DownloadRequest(BaseModel):
    sections: Dict[str, str]
    format: str  
    filename: Optional[str] = "facilities_draft"  
    remove_citations: Optional[bool] = False 


def generate_facilities_pdf(sections: Dict[str, str], filename: str = "facilities_draft", remove_citations: bool = False) -> bytes:
    """Generate PDF from facilities sections """
    from reportlab.lib.units import inch
    from reportlab.platypus import PageTemplate, Frame
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus.flowables import PageBreak
    import re
    
    buffer = BytesIO()
    
    def convert_markdown_to_html(text):
        """Convert markdown formatting to ReportLab HTML-like tags"""

        text = text.replace('&', '&amp;')
        
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        
        text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
        

        text = text.replace('<b>', '<<<BOLD_OPEN>>>')
        text = text.replace('</b>', '<<<BOLD_CLOSE>>>')
        text = text.replace('<i>', '<<<ITALIC_OPEN>>>')
        text = text.replace('</i>', '<<<ITALIC_CLOSE>>>')
        
        text = text.replace('<', '&lt;').replace('>', '&gt;')
        
        text = text.replace('<<<BOLD_OPEN>>>', '<b>')
        text = text.replace('<<<BOLD_CLOSE>>>', '</b>')
        text = text.replace('<<<ITALIC_OPEN>>>', '<i>')
        text = text.replace('<<<ITALIC_CLOSE>>>', '</i>')
        
        return text
    
    def add_footer(canvas, doc):
        canvas.saveState()
        
        page_num = canvas.getPageNumber()
        
        footer_left = filename
        footer_right = f"Page {page_num}"
        
        canvas.setFont("Helvetica", 9)
        canvas.drawString(0.5*inch, 0.5*inch, footer_left)  # Left side
        canvas.drawRightString(letter[0] - 0.5*inch, 0.5*inch, footer_right)  # Right side
        
        canvas.restoreState()
    
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=1*inch  # Extra space for footer
    )
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20
    )
    subsection_style = ParagraphStyle(
        'SubsectionHeading',
        parent=styles['Heading3'],
        fontSize=10,
        spaceAfter=8,
        spaceBefore=16
    )
    normal_style = styles['Normal']
    
    # Add title
    story.append(Paragraph("Pilot GenAI Generated Draft", title_style))
    story.append(Spacer(1, 20))
    
    # Add content
    for section_label, section_text in sections.items():
        if section_text and section_text.strip():
            story.append(Paragraph(section_label, heading_style))
            

            processed_text = section_text
            if remove_citations:
                import re

                processed_text = re.sub(r'\[[^\]]*\.pdf[^\]]*\]', '', processed_text)
                processed_text = re.sub(r'\[https?://[^\]]+\]', '', processed_text)

                processed_text = re.sub(r'[ \t]+', ' ', processed_text)  # Replace multiple spaces/tabs with single space
                processed_text = re.sub(r'\n\s*\n', '\n\n', processed_text)  # Preserve paragraph breaks
                processed_text = processed_text.strip()
            

            lines = processed_text.split('\n')
            logging.info(f"PDF Processing: Processing {len(lines)} lines for section {section_label}")
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph:

                        paragraph_text = ' '.join(current_paragraph)
                        formatted_text = convert_markdown_to_html(paragraph_text)
                        story.append(Paragraph(formatted_text, normal_style))
                        story.append(Spacer(1, 6))  
                        current_paragraph = []
                    continue
                
                is_subsection_heading = (
                    line.startswith('### ') or 
                    (line.endswith(':') and 
                     any(keyword in line for keyword in ['Core Instrumentation', 'Computing and Data Resources', 
                                                        'Internal Facilities', 'External Facilities', 'Special Infrastructure']))
                )
                

                if line.startswith('### '):
                    logging.info(f"PDF Processing: Found markdown heading: '{line}'")
                
                if is_subsection_heading:

                    if current_paragraph:
                        paragraph_text = ' '.join(current_paragraph)
                        formatted_text = convert_markdown_to_html(paragraph_text)
                        story.append(Paragraph(formatted_text, normal_style))
                        current_paragraph = []
                    
                    clean_heading = line.replace('### ', '')
                    formatted_heading = convert_markdown_to_html(clean_heading)
                    logging.info(f"PDF Processing: Adding subsection heading: '{clean_heading}' with subsection_style")
                    story.append(Paragraph(formatted_heading, subsection_style))
                    story.append(Spacer(1, 6))  # Add small space after subsection heading
                else:

                    if (line.startswith('###') or 
                        (line.endswith(':') and any(keyword in line for keyword in ['Core Instrumentation', 'Computing and Data Resources', 
                                                                                   'Internal Facilities', 'External Facilities', 'Special Infrastructure']))):
                        logging.info(f"PDF Processing: Found subsection heading in regular content: '{line}'")

                        if current_paragraph:
                            paragraph_text = ' '.join(current_paragraph)
                            formatted_text = convert_markdown_to_html(paragraph_text)
                            story.append(Paragraph(formatted_text, normal_style))
                            story.append(Spacer(1, 6))
                            current_paragraph = []
                        
                        clean_heading = line.replace('### ', '')
                        formatted_heading = convert_markdown_to_html(clean_heading)
                        story.append(Paragraph(formatted_heading, subsection_style))
                        story.append(Spacer(1, 6))
                    else:
                        current_paragraph.append(line)
            

            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                formatted_text = convert_markdown_to_html(paragraph_text)
                story.append(Paragraph(formatted_text, normal_style))
                story.append(Spacer(1, 6))  
            
            story.append(Spacer(1, 12))
    
    # First pass: count pages by building to a temporary buffer
    import copy
    temp_buffer = BytesIO()
    temp_doc = SimpleDocTemplate(
        temp_buffer, 
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=1*inch
    )
    
    # Create a deep copy of the story for counting
    story_copy = copy.deepcopy(story)
    temp_doc.build(story_copy)
    total_pages = temp_doc.page
    
    # Update the footer function with the correct total
    def add_footer_with_total(canvas, doc):
        canvas.saveState()
        
        page_num = canvas.getPageNumber()
        
        footer_left = filename
        footer_right = f"Page {page_num} of {total_pages}"
        
        canvas.setFont("Helvetica", 9)
        canvas.drawString(0.5*inch, 0.5*inch, footer_left)  # Left side
        canvas.drawRightString(letter[0] - 0.5*inch, 0.5*inch, footer_right)  # Right side
        
        canvas.restoreState()
    
    # Second pass: build the actual document with correct page numbers
    doc.build(story, onFirstPage=add_footer_with_total, onLaterPages=add_footer_with_total)
    buffer.seek(0)
    return buffer.getvalue()


def generate_facilities_doc(sections: Dict[str, str], filename: str = "facilities_draft", remove_citations: bool = False) -> bytes:
    """Generate DOCX from facilities sections"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Pilot GenAI Generated Draft', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    
    # Add spacing after title
    doc.add_paragraph()
    
    def remove_citations_from_text(text: str) -> str:
        """Remove citations from text"""
        if not remove_citations:
            return text
        
        # Remove PDF citations
        text = re.sub(r'\[[^\]]*\.pdf[^\]]*\]', '', text)
        # Remove web URL citations
        text = re.sub(r'\[https?://[^\]]+\]', '', text)
        # Clean up extra spaces
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()
    
    def add_formatted_paragraph(doc, text: str, is_heading: bool = False, is_subsection: bool = False):
        """Add a paragraph with proper formatting"""
        if is_heading:
            para = doc.add_heading(text, level=1)
            para_run = para.runs[0]
            para_run.font.size = Pt(14)
            para_run.font.bold = True
            para_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        elif is_subsection:
            para = doc.add_heading(text, level=2)
            para_run = para.runs[0]
            para_run.font.size = Pt(12)
            para_run.font.bold = True
            para_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        else:
            para = doc.add_paragraph()
            # Process text for bold/italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    # Italic text
                    run = para.add_run(part[1:-1])
                    run.italic = True
                elif part.strip():
                    # Regular text
                    para.add_run(part)
        
        return para
    
    # Process each section
    for section_label, section_text in sections.items():
        if section_text and section_text.strip():
            # Add section heading
            add_formatted_paragraph(doc, section_label, is_heading=True)
            
            # Process section text
            processed_text = remove_citations_from_text(section_text)
            
            # Split by lines to handle subsections
            lines = processed_text.split('\n')
            current_paragraph_text = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph_text:
                        para_text = ' '.join(current_paragraph_text)
                        add_formatted_paragraph(doc, para_text)
                        current_paragraph_text = []
                    doc.add_paragraph()  # Add blank line
                    continue
                
                # Check if it's a subsection heading
                is_subsection_heading = (
                    line.startswith('### ') or 
                    (line.endswith(':') and 
                     any(keyword in line for keyword in ['Core Instrumentation', 'Computing and Data Resources', 
                                                        'Internal Facilities', 'External Facilities', 'Special Infrastructure']))
                )
                
                if is_subsection_heading:
                    # Finish current paragraph
                    if current_paragraph_text:
                        para_text = ' '.join(current_paragraph_text)
                        add_formatted_paragraph(doc, para_text)
                        current_paragraph_text = []
                    
                    # Add subsection heading
                    clean_heading = line.replace('### ', '')
                    add_formatted_paragraph(doc, clean_heading, is_subsection=True)
                else:
                    current_paragraph_text.append(line)
            
            # Add remaining paragraph
            if current_paragraph_text:
                para_text = ' '.join(current_paragraph_text)
                add_formatted_paragraph(doc, para_text)
            
            # Add spacing after section
            doc.add_paragraph()
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@router.post("/download")
async def download_facilities_document(
    request: Request,
    download_data: DownloadRequest,
    user=Depends(get_verified_user)
):
    """
    Download facilities document as PDF or DOC
    """
    if not request.app.state.config.ENABLE_FACILITIES.get(user.email):
        raise HTTPException(status_code=403, detail="Facilities feature is not enabled for this user")
    
    try:
        if download_data.format not in ["pdf", "doc"]:
            raise HTTPException(status_code=400, detail="Invalid format. Only 'pdf' and 'doc' are supported")
        
        if not download_data.sections:
            raise HTTPException(status_code=400, detail="No sections provided")
        
        if download_data.format == "pdf":
            content = generate_facilities_pdf(download_data.sections, download_data.filename, download_data.remove_citations)
            media_type = "application/pdf"
            filename = f"{download_data.filename}.pdf"
        else:  # doc format
            # DOC format always removes citations (as per requirement)
            content = generate_facilities_doc(download_data.sections, download_data.filename, remove_citations=True)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{download_data.filename}.docx"
        
        return Response(
            content=content,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Error generating {download_data.format} document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate document: {str(e)}")

# Prompt template for extracting form data from files
EXTRACT_FORM_DATA_PROMPT = """You are an expert at analyzing research documents and extracting relevant information for grant proposal forms.

Analyze the provided document content and extract information for the following section:

**Section: {section_label}**
{section_instructions}

**Document Content:**
\"\"\"
{document_content}
\"\"\"

**CRITICAL EXTRACTION REQUIREMENTS:**
1. Extract ONLY information that is directly relevant to this specific section based on the definition provided
2. Focus on MAIN KEYWORDS and KEY INFORMATION - identify the most important facts
3. Maximum 5 lines of text - be concise but comprehensive
4. Include all main points that fit the section definition
5. If the section asks for specific details (like square footage, location, equipment names, models), extract those if mentioned
6. **CRITICAL: If no relevant information is found for this section, you MUST NOT return anything (nothing, not even "empty string" or "not available" or any other text)**
7. Do not invent or assume information not present in the document
8. Focus on factual information that can be directly used in a grant proposal
9. Each line should contain distinct, important information - avoid repetition
10. Prioritize specific details (names, numbers, locations) over general statements

**Output Format:**
- Maximum 5 lines
- Each line should be a complete, meaningful statement
- Include key keywords, names, numbers, and specific details
- Cover all main aspects mentioned in the section definition
- **If nothing is found: return absolutely nothing (empty string, no text at all - do NOT write "empty string", "not available", "no information found", or any other placeholder text)**

Return ONLY the extracted information for this section (max 5 lines), nothing else. If no relevant information is found, return a completely empty string with no text whatsoever."""

@router.post("/extract-form-data", response_model=ExtractFormDataResponse)
async def extract_form_data_from_files(request: Request, form_data: ExtractFormDataRequest, user=Depends(get_verified_user)):
    """
    Extract form data from uploaded files using LLM
    Only works when facilities is enabled for the user
    """
    if not request.app.state.config.ENABLE_FACILITIES.get(user.email):
        raise HTTPException(status_code=403, detail="Facilities feature is not enabled for this user")
    
    try:
        if form_data.sponsor not in ["NSF", "NIH"]:
            raise HTTPException(status_code=400, detail="Invalid sponsor. Must be 'NSF' or 'NIH'")
        
        if not form_data.files or len(form_data.files) == 0:
            raise HTTPException(status_code=400, detail="No files provided")
        
        # Get model information
        model = request.app.state.MODELS.get(form_data.model)
        if not model:
            raise HTTPException(status_code=400, detail="Model not found")
        
        # Process attached files to get their content
        attached_files_sources = []
        try:
            from open_webui.utils.middleware import chat_completion_files_handler
            
            mock_body = {
                "model": form_data.model,
                "metadata": {
                    "files": form_data.files
                },
                "messages": [{"role": "user", "content": "Extract form data from attached files"}]
            }
            
            processed_body, file_flags = await chat_completion_files_handler(request, mock_body, user)
            attached_files_sources = file_flags.get("sources", [])
            
            logging.info(f"Processed {len(attached_files_sources)} file sources for extraction")
            
        except Exception as e:
            logging.error(f"Error processing attached files: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to process files: {str(e)}")
        
        # Combine all document content from files
        all_document_content = []
        for source in attached_files_sources:
            documents = source.get("document", [])
            for doc in documents:
                if isinstance(doc, str):
                    all_document_content.append(doc)
        
        combined_content = "\n\n".join(all_document_content)
        
        if not combined_content.strip():
            raise HTTPException(status_code=400, detail="No content extracted from files")
        
        logging.info(f"Combined document content length: {len(combined_content)} characters")
        
        # section instructions based on sponsor with clear, detailed definitions
        # NSF section definitions
        nsf_section_definitions = {
        "projectTitle": {
            "label": "1. Project Title",
            "instructions": """Extract the formal research project title.

    DEFINITION:
    The official title of the NSF proposal or research project.

    Look for (in order of preference):
    - The title on the NSF cover sheet or first page.
    - The title at the top of the "Project Summary" or "Project Overview" page.

    Guidelines:
    - Extract the exact title text as written (do not paraphrase).
    - Exclude labels like 'Project Summary', 'Overview', 'Intellectual Merit', etc.
    - If multiple titles appear (e.g. main + internal code), choose the main NSF project title."""
        },

        "researchSpaceFacilities": {
            "label": "2. Research Space and Facilities",
            "instructions": """Extract information about physical research spaces and facilities.

    DEFINITION (FACILITIES):
    Describe relevant Laboratory, Clinical, Animal, Computer, and Office facilities/resources that are
    physically available to the project. Focus on physical spaces, not individual pieces of equipment.

    Look for:
    - Laboratory spaces (research labs, experimental labs, wet/dry labs, structural labs, robotics labs).
    - Clinical or animal facilities, if present.
    - Computer facilities as physical spaces (data centers, server rooms, visualization labs).
    - Office facilities (research offices, student workspaces, shared collaboration areas).
    - Large-scale test spaces (e.g., high-bay labs, large-scale testing halls, field test sites, testbeds as physical locations).
    - Location details (building names, centers, campuses, addresses).
    - Physical characteristics (e.g., high-bay, reaction floor, cleanroom, number of floors, capacity).
    - Any mention of floor area / square footage (if present), but DO NOT extract any costs.

    Guidelines:
    - Focus on describing the physical spaces and their suitability for the proposed research.
    - Do NOT include detailed equipment lists here (those go under Core Instrumentation or Equipment).
    - Do NOT extract any financial or budget information."""
        },

        "coreInstrumentation": {
            "label": "3. Core Instrumentation",
            "instructions": """Extract information about specialized research instruments and experimental platforms.

    DEFINITION:
    Specialized scientific and engineering equipment used to conduct experiments or measurements.
    These are typically non-general-purpose instruments or platforms.

    Include (examples, not exhaustive):
    - Wireless and networking testbeds (e.g., mmWave testbeds, COSMOS nodes, UAV communication testbeds).
    - Software-defined radios (SDRs), RF frontends, antennas, transceivers.
    - National Instruments / LabVIEW PXI systems, custom RF chains, measurement racks.
    - Robotic platforms (humanoids, quadrupeds, industrial arms) used for research.
    - Structural testing rigs, load frames, reaction frames, actuators, sensing systems.
    - Scientific instruments, sensors, motion capture systems, specialized cameras, etc.
    - Any clearly experimental hardware platform built or used specifically for research.

    Look for:
    - Named testbeds, platforms, or systems explicitly used in experiments.
    - Descriptions of what the instrument or platform enables (e.g., “high-data-rate mmWave experimentation”, “full-scale building connection testing”).
    - Distinction between experimental hardware and general-purpose computing.

    DO NOT include:
    - General-purpose computers, servers, HPC clusters, or storage systems (these go under Computing and Data Resources).
    - Purely physical building/space descriptions (those go under Research Space and Facilities).

    Guidelines:
    - Focus on capabilities and relevance to the project, not price or purchase details.
    - Do NOT extract any financial or budget information."""
        },

        "computingDataResources": {
            "label": "4. Computing and Data Resources",
            "instructions": """Extract information about computational and data infrastructure.

    DEFINITION:
    Any computational, storage, or data-related infrastructure that supports simulation, data analysis,
    modeling, visualization, or large-scale experiments.

    Include:
    - High Performance Computing (HPC) clusters and supercomputers.
    - University research computing facilities (e.g., central research computing, research data centers).
    - GPU clusters, parallel computing systems, virtualization environments, cloud or edge resources.
    - Specialized simulator/emulator resources (e.g., NS3 setups, network emulation environments) when described as computing infrastructure.
    - Storage systems and data infrastructure (SAN/NAS, research data storage, archival systems).
    - Secure data enclaves, big-data platforms (Hadoop, Spark clusters, large databases).
    - Visualization facilities that are primarily computing-centered (e.g., tiled display walls with dedicated compute).

    Look for:
    - Named clusters or facilities (e.g., 'Shamu cluster', 'Research Computing Facility', 'HPC cluster').
    - Descriptions of CPU cores, memory, TB of storage, networking, or similar capacity-related info.
    - References to secure computing environments for sensitive or large datasets.

    DO NOT include:
    - Experimental instruments (SDRs, testbeds, robots, sensors) — those belong to Core Instrumentation.
    - Purely physical descriptions of buildings or rooms (those belong to Research Space and Facilities).

    Guidelines:
    - Focus on the computing and data capabilities provided to the project.
    - Do NOT extract any financial or budget information (no costs, prices, or funding amounts)."""
        },

        "internalFacilitiesNYU": {
            "label": "5a. Internal Facilities (NYU)",
            "instructions": """Extract facilities, resources, and infrastructure explicitly provided by NYU (New York University).

    DEFINITION:
    Facilities, resources, or infrastructure located at or provided by NYU.

    Look for:
    - Mentions of NYU, New York University, internal university facilities
    - NYU-specific labs, centers, institutes, facilities
    - NYU resources, NYU infrastructure, NYU equipment
    - Facilities located at NYU campuses or NYU buildings
    - NYU-provided services, resources, or support

    Guidelines:
    - This section is an NYU-specific view: re-group NYU-related items already captured in other sections and summarize them as internal resources.
    - Only include resources clearly identified as part of NYU.
    - Do NOT include facilities or equipment that are clearly external or at other institutions.
    - Do NOT extract any financial or budget information."""
        },

        "externalFacilitiesOther": {
            "label": "5b. External Facilities (Other Institutions)",
            "instructions": """Extract facilities, resources, and infrastructure at institutions other than NYU.

    DEFINITION:
    Resources provided or made available by collaborating universities, research centers, industry partners,
    or national facilities that are not part of NYU.

    Look for:
    - Named partner universities or institutions and their labs, centers, or facilities.
    - External structural labs, testbeds, robotics facilities, VR/AR labs, or large-scale testing sites.
    - External HPC or computing resources (e.g., partner institution clusters, national centers) explicitly available to the project.
    - Shared experimental platforms or testbeds hosted at other institutions.

    Guidelines:
    - Only include facilities clearly attributed to external institutions (non-NYU).
    - Summarize these as external resources that complement NYU’s internal facilities.
    - Do NOT include NYU facilities in this section."""
        },

        "specialInfrastructure": {
            "label": "6. Special Infrastructure",
            "instructions": """Extract information about unique, one-of-a-kind, or particularly notable infrastructure.

    DEFINITION:
    Infrastructure that is special, unusual, or uniquely important to the proposed project and does not fit
    cleanly into a single category above, or is called out in the proposal as a distinctive capability.

    Look for:
    - National or regional testbeds or platforms (e.g., large-scale wireless testbeds, full-scale building testbeds).
    - One-of-a-kind experimental systems or prototypes (e.g., custom multi-gigabit mesh backhaul systems, unique SDR platforms).
    - Specialized infrastructure with strict environmental or operational requirements (e.g., cleanrooms, vibration-isolated labs, specialized VR/AR walls).
    - Large integrated systems that combine space, equipment, and computing into a distinctive resource.

    Guidelines:
    - Focus on what makes the infrastructure special or unique relative to standard labs/equipment.
    - This section can reference facilities already mentioned but highlight their uniqueness and role in the project.
    - Do NOT extract any financial or budget information."""
        },

        "equipment": {
            "label": "7. Equipment",
            "instructions": """Extract information about equipment, instruments, and tools that are not already clearly captured
    under Core Instrumentation, or that are listed more generally.

    DEFINITION:
    Research equipment, instruments, and tools (broadly defined), including both specialized and
    supporting devices.

    Include:
    - Named equipment, instruments, and tools with models or specifications (if mentioned).
    - Laboratory tools, measurement devices, and supporting hardware that may be secondary to the main core instrumentation.
    - General equipment that may not be explicitly part of a named testbed but still supports the research.

    Relationship to Core Instrumentation:
    - Core Instrumentation should capture the main experimental platforms and major research instruments.
    - Equipment here can include additional or supporting instruments, tools, or devices, especially when listed in more generic equipment sections.

    Look for:
    - Equipment names, types, models, or brief specs (if present).
    - Descriptions of capabilities or roles in the research.
    - Indications of availability and usage for the project.

    Guidelines:
    - Do NOT duplicate long lists already captured under Core Instrumentation unless needed to complete the picture.
    - Do NOT include computers, servers, or HPC clusters here (those belong to Computing and Data Resources).
    - Do NOT extract any financial or budget information (no prices, no purchase details)."""
        }
    }

        # NIH section definitions
        nih_section_definitions = {
        "projectTitle": {
            "label": "1. Project Title",
            "instructions": """Extract the formal research project title.

    DEFINITION:
    The official title of the NIH proposal or research project.

    Look for (in order of preference):
    - The title on the NIH cover sheet, application face page, or first page.
    - The title at the top of the "Project Summary/Abstract" or "Specific Aims" page.

    Guidelines:
    - Extract the exact title text as written (do not paraphrase).
    - Exclude labels like 'Project Summary', 'Specific Aims', 'Abstract', etc.
    - If multiple titles appear (e.g. main + internal code), choose the main NIH project title."""
        },

        "laboratory": {
            "label": "2. Laboratory",
            "instructions": """Extract information about laboratory spaces where the research will be performed.

    DEFINITION:
    Physical laboratory spaces used for the proposed research, including their location, size, and outfitting.

    Look for:
    - Lab location: campus, building name, and room number.
    - Whether the PI has a dedicated laboratory; list its size in square feet if mentioned.
    - If the PI shares space with another PI or works in a mentor's laboratory, note that and explain what space is available for the project.
    - Lab outfitting and capabilities pertinent to the project: biological safety cabinets, chemical fume hoods, tissue culture incubators, bench- and micro-centrifuges, refrigerators, freezers, and other relevant equipment.
    - If sharing space, what resources are specifically available for the study.
    - If more than one laboratory will be used, extract information for each lab separately.

    Guidelines:
    - Focus on where the work will be performed and how the lab is equipped.
    - Include specific capabilities pertinent to the proposed research.
    - Do NOT include animal facilities here (those go under Animal).
    - Do NOT include computers/servers here (those go under Computer).
    - Do NOT extract any financial or budget information."""
        },

        "animal": {
            "label": "3. Animal",
            "instructions": """Extract information about animal facilities and resources for the study.

    DEFINITION:
    Animal housing, procedure rooms, surgical facilities, and institutional animal care resources
    relevant to the proposed research.

    Look for:
    - Where animals are housed and proximity to the PI's laboratory.
    - Specific procedure rooms or surgical suites available for animal studies.
    - Equipment available for animal research (surgical equipment, imaging for animals, behavioral testing apparatus).
    - Institutional resources: veterinary care, basic husbandry, IACUC support.
    - Animal facility certifications or accreditations (e.g., AAALAC).

    Guidelines:
    - Only include information if the study involves animals.
    - Focus on housing location, proximity to lab, available procedure rooms, equipment, and institutional support.
    - Do NOT extract any financial or budget information."""
        },

        "computer": {
            "label": "4. Computer",
            "instructions": """Extract information about computer resources available for the research.

    DEFINITION:
    Computer hardware, software, networking, and computing infrastructure available to the PI and
    research staff for the proposed project.

    Look for:
    - PCs, workstations, and their operating systems.
    - Basic software needed for the research (statistical software, graphical software, Office suite, specialized analysis tools).
    - Internet access and network connectivity.
    - Computers available to lab staff, students, or research personnel.
    - Specialized computing resources: mainframe access, HPC clusters, GPU resources, cloud computing.
    - Data storage and backup systems.

    Guidelines:
    - Include both personal computing resources and shared/institutional computing infrastructure.
    - Describe availability and accessibility of computing resources.
    - Do NOT extract any financial or budget information."""
        },

        "office": {
            "label": "5. Office",
            "instructions": """Extract information about office spaces for the PI and research personnel.

    DEFINITION:
    Dedicated office space for the PI, research staff, students, and other project personnel.

    Look for:
    - Dedicated office space for the PI: size in square feet, location (campus, building, room number).
    - Proximity of office to the laboratory.
    - Office space available to employees, students, postdocs, and other research personnel.
    - Whether office space is dedicated or shared, and its size.
    - Whether office space is located in or near the lab.

    Guidelines:
    - Focus on office locations, sizes, and who has access.
    - Include proximity information relative to the lab.
    - Do NOT extract any financial or budget information."""
        },

        "clinical": {
            "label": "6. Clinical",
            "instructions": """Extract information about clinical resources available for the study.

    DEFINITION:
    Clinical facilities, resources, and support available to the PI for clinical research activities.

    Look for:
    - Clinical resources available to support the research (clinical trial units, patient recruitment facilities, clinical labs).
    - Support from Clinical and Translational Science Institutes (CTSI) or similar programs (e.g., ACTSI).
    - Clinical cores: biostatistics cores, biorepositories, imaging cores, clinical pharmacology units.
    - Patient populations or clinical cohorts available for the study.
    - Clinical data infrastructure (electronic health records access, clinical databases).

    Guidelines:
    - Only include information if the study has a clinical component.
    - Focus on clinical resources, institutional support, and available clinical infrastructure.
    - Do NOT extract any financial or budget information."""
        },

        "other": {
            "label": "7. Other",
            "instructions": """Extract information about other institutional, departmental, or divisional resources.

    DEFINITION:
    Any other equipment, resources, or institutional support necessary for the project that does not fit
    into the Laboratory, Animal, Computer, Office, or Clinical categories.

    Look for:
    - Institutional resources: libraries, shared core facilities, fabrication shops, machine shops.
    - Departmental or divisional resources and support services.
    - Collaborative arrangements or shared resources with other departments or institutions.
    - Training and mentorship resources, especially for Early Stage Investigators (ESIs).
    - Institutional investment in the investigator's success: resources for classes, travel, training.
    - Collegial support: career enrichment programs, peer groups, supervision assistance.
    - Administrative and logistical support: management, oversight, best practices training.
    - Protected time for research with salary support.
    - Special facilities for working with biohazards or potentially dangerous substances.
    - Any unique features of the scientific environment that contribute to probability of success.

    Guidelines:
    - This is a catch-all section for resources not covered by other NIH sections.
    - Include institutional support and scientific environment contributions.
    - For ESIs, include institutional investment details if available.
    - Do NOT duplicate information already captured in other sections.
    - Do NOT extract any financial or budget information."""
        },

        "equipment": {
            "label": "8. Equipment",
            "instructions": """Extract information about major equipment available for the research.

    DEFINITION:
    Research equipment, instruments, and tools available for the proposed project, including both
    specialized and general-purpose research devices.

    Look for:
    - Named equipment, instruments, and tools with models or specifications (if mentioned).
    - Major research instruments: microscopes, spectrometers, sequencers, imaging systems, flow cytometers.
    - Laboratory equipment: centrifuges, PCR machines, gel electrophoresis systems, mass spectrometers.
    - Shared or core facility equipment available for use.
    - Any equipment specifically relevant to the proposed research methods.

    Guidelines:
    - Focus on equipment that is directly relevant to the proposed research.
    - Include equipment names, types, models, and brief specs if present.
    - Do NOT include computers or general office equipment (those go under Computer or Office).
    - Do NOT extract any financial or budget information (no prices, no purchase details)."""
        }
    }

        # Select section definitions based on sponsor
        section_definitions = nsf_section_definitions if form_data.sponsor == "NSF" else nih_section_definitions

        # sections to extract based on sponsor
        if form_data.sponsor == "NSF":
            sections_to_extract = ["projectTitle", "researchSpaceFacilities", "coreInstrumentation", 
                                  "computingDataResources", "internalFacilitiesNYU", "externalFacilitiesOther", 
                                  "specialInfrastructure"]
        else:  # NIH
            sections_to_extract = ["projectTitle", "laboratory", "animal", "computer", 
                                  "office", "clinical", "other", "equipment"]
        
        extracted_form_data = {}
        
        for section_key in sections_to_extract:
            section_info = section_definitions[section_key]
            
            prompt = EXTRACT_FORM_DATA_PROMPT.format(
                section_label=section_info["label"],
                section_instructions=section_info["instructions"],
                document_content=combined_content[:50000]  # Limit content to avoid token limits
            )
            
            try:
                extracted_text = await call_llm_direct(prompt, form_data.model, user, request)
                extracted_text = extracted_text.strip()
                
                if extracted_text and len(extracted_text) > 10:
                    extracted_form_data[section_key] = extracted_text
                    logging.info(f"Extracted {len(extracted_text)} characters for {section_key}")
                else:
                    extracted_form_data[section_key] = ""
                    logging.info(f"No relevant content found for {section_key}")
                    
            except Exception as e:
                logging.error(f"Error extracting data for {section_key}: {e}")
                extracted_form_data[section_key] = ""
        
        return ExtractFormDataResponse(
            success=True,
            form_data=extracted_form_data,
            error=None
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in extract_form_data_from_files: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
